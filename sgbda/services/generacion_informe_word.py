import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.conf import settings
from docx import Document
from django.template.defaultfilters import filesizeformat

# ============================================================
# CREACION DE DOCUMENTO WORD
# ============================================================
def generar_reporte_oracle(datos, buffer=None, ruta_salida="/tmp/reporte_oracle.docx",
                           empresa=None, conexion=None):
    """
    Genera un reporte Word (.docx) con las métricas de Oracle.
    
    Args:
        datos_json: dict con la estructura del JSON de métricas
        buffer: BytesIO opcional. Si se pasa, guarda ahí (para descarga web).
        ruta_salida: str opcional. Ruta del archivo si no se pasa buffer.
    """
    # ============================================================
    # ABRIR PLANTILLA
    # ============================================================
    if not os.path.exists(settings.PLANTILLA_WORD_PATH):
        raise FileNotFoundError(f"No existe la plantilla: {settings.PLANTILLA_WORD_PATH}")
    
    doc = Document(settings.PLANTILLA_WORD_PATH)
    
    # ============================================================
    # CONFIGURAR PÁGINA
    # ============================================================
    section = doc.sections[0]
    section.page_height = Cm(27.94)
    section.page_width = Cm(21.59)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # ============================================================
    # FUNCIONES AUXILIARES
    # ============================================================
    def parrafo_vacio(espacio_pt=12):
        p = doc.add_paragraph()
        p.space_after = Pt(espacio_pt)
        return p



    def texto_posicionado(doc, texto, posicion='centro', tamano=11, bold=False,
                        color=(64,64,64), fuente='Calibri'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ─────────────────────────────────────────
        # ESPACIO ARRIBA (antes del texto)
        # Aumentá o disminuí este número para subir/bajar el texto
        # ─────────────────────────────────────────
        p.paragraph_format.space_before = Pt(0)   # ← 0 = sin espacio arriba

        # ─────────────────────────────────────────
        # TEXTO
        # ─────────────────────────────────────────
        run = p.add_run(texto)
        run.font.size = Pt(tamano)
        run.font.name = fuente
        run.bold = bold
        run.font.color.rgb = RGBColor(*color)

        # ─────────────────────────────────────────
        # ESPACIO ABAJO (después del texto)
        # Aumentá o disminuí este número para separar más/menos del siguiente elemento
        # ─────────────────────────────────────────
        p.paragraph_format.space_after = Pt(0)    # ← 0 = sin espacio abajo

        return p

    def texto_centrado(texto, tamano=11, bold=True, color=(64,64,64), fuente='Calibri', espacio_despues=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.font.size = Pt(tamano)
        run.font.name = fuente
        run.bold = bold
        run.font.color.rgb = RGBColor(*color)
        p.space_after = Pt(espacio_despues)
        return p

    def tabla_sin_bordes(filas, col_widths, alineacion_col1=WD_ALIGN_PARAGRAPH.RIGHT):
        """filas: lista de [texto_col1, texto_col2]"""
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False
        table.allow_autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Quitar bordes globales de la tabla
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        for fila in filas:
            row = table.add_row().cells
            row[0].text = fila[0]
            row[1].text = fila[1]
            
            row[0].width = col_widths[0]
            row[1].width = col_widths[1]
            
            # Formato col 1 (título)
            for paragraph in row[0].paragraphs:
                paragraph.alignment = alineacion_col1
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(80, 80, 80)
            
            # Formato col 2 (valor)
            for paragraph in row[1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(64, 64, 64)
            
            # Quitar bordes de celdas individualmente
            for cell in row:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
        
        return table

    def agregar_titulo(texto, nivel=1, centrado=False):
        p = doc.add_heading(texto, level=nivel)
        if centrado:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def formatear_celda_contenido(cell, texto, alineacion_horizontal=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = texto
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        for paragraph in cell.paragraphs:
            paragraph.alignment = alineacion_horizontal
            # Espaciado simétrico para centrar verticalmente el texto
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(64, 64, 64)

    def agregar_tabla_contenidos(doc, items):
        """
        items: lista de tuplas [(titulo, numero_pagina), ...]
        Ejemplo: [("1. Resumen Ejecutivo", "3"), ("2. Estado del Servidor", "3")]
        """
        # Separación 2: otro párrafo vacío (Word nunca colapsa dos seguidos con contenido)
        # sep2 = doc.add_paragraph()
        # sep2.add_run(' ')
        # sep2.paragraph_format.space_after = Pt(4)

        # Título "Contenido"
        p = doc.add_paragraph()
        run = p.add_run("Contenido")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)

        # Separación 1: párrafo vacío con run invisible
        sep1 = doc.add_paragraph()
        sep1.add_run(' ')
        sep1.paragraph_format.space_after = Pt(10)

        # Tabla sin bordes visibles
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Ancho total de la tabla (ajusta a tus márgenes)
        table.width = Cm(17)
        
        for titulo, pagina in items:
            row = table.add_row().cells
            
            row[0].width = Cm(14)
            row[1].width = Cm(3)
            
            # Formatear celdas con centrado vertical
            formatear_celda_contenido(row[0], titulo, WD_ALIGN_PARAGRAPH.LEFT)
            formatear_celda_contenido(row[1], str(pagina), WD_ALIGN_PARAGRAPH.RIGHT)
            
            # Altura de fila
            tr = row[0]._tc.getparent()
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '500')
            trHeight.set(qn('w:hRule'), 'atLeast')
            trPr.append(trHeight)
            
            # Bordes solo abajo
            for cell in row:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    border.set(qn('w:sz'), '0')
                    tcBorders.append(border)
                
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:color'), 'BFBFBF')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)

        doc.add_paragraph()
    
    def agregar_parrafo(texto, negrita=False, centrado=False, tamano=11, color=None):
        p = doc.add_paragraph()
        run = p.add_run(texto)
        run.font.size = Pt(tamano)
        run.font.name = 'Calibri'
        run.bold = negrita
        if color:
            run.font.color.rgb = RGBColor(*color)
        if centrado:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p
    
    def agregar_tabla(headers, filas, ancho_total=Cm(17)):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        table.width = ancho_total
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
            # Fondo azul
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '9F2B20')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Filas
        for fila in filas:
            row_cells = table.add_row().cells
            for i, valor in enumerate(fila):
                row_cells[i].text = str(valor)
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        return table

    def tabla_tipo_2(headers, filas, ancho_total=Cm(17), mostrar_encabezado=True):
        num_cols = len(headers)
        
        if mostrar_encabezado:
            table = doc.add_table(rows=1, cols=num_cols)
        else:
            table = doc.add_table(rows=0, cols=num_cols)
        
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        table.width = ancho_total

        # Definir anchos: 
        ancho_col1 = Cm(5.24)
        
        columnas_fijas = {0: ancho_col1}
        columnas_restantes = num_cols - len(columnas_fijas)
        
        if columnas_restantes > 0:
            ancho_usado = sum(columnas_fijas.values())
            ancho_resto = int((ancho_total - ancho_usado) / columnas_restantes)
        
        anchos = []
        for i in range(num_cols):
            if i in columnas_fijas:
                anchos.append(columnas_fijas[i])
            else:
                anchos.append(ancho_resto)
        
        def aplicar_anchos(row):
            for i, cell in enumerate(row.cells):
                cell.width = anchos[i]
        
        for i, width in enumerate(anchos):
            table.columns[i].width = width
                
        def aplicar_estilo(cell):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '9F2B20')
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Encabezados (opcional)
        if mostrar_encabezado:
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                aplicar_estilo(hdr_cells[i])
            aplicar_anchos(table.rows[0])

        
        # Filas
        for fila in filas:
            row_cells = table.add_row().cells
            for i, valor in enumerate(fila):
                row_cells[i].text = str(valor)
                for paragraph in row_cells[i].paragraphs:
                    if i == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        aplicar_estilo(row_cells[i])
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            aplicar_anchos(table.rows[-1])
        
        return table

    def table_for_tablespace(headers, filas, ancho_total=Cm(17)):
        num_cols = len(headers)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        table.width = ancho_total
        
        # Definir anchos: 
        ancho_col1 = Cm(4.5)
        ancho_col2 = Cm(4.1)
        ancho_col3 = Cm(2)
        ancho_col6 = Cm(1.7)  
        ancho_col7 = Cm(1.75)

        
        columnas_fijas = {0: ancho_col1, 1: ancho_col2, 2: ancho_col3, 6: ancho_col7, 5: ancho_col6}
        columnas_restantes = num_cols - len(columnas_fijas)
        
        if columnas_restantes > 0:
            ancho_usado = sum(columnas_fijas.values())
            ancho_resto = int((ancho_total - ancho_usado) / columnas_restantes)
        
        anchos = []
        for i in range(num_cols):
            if i in columnas_fijas:
                anchos.append(columnas_fijas[i])
            else:
                anchos.append(ancho_resto)
        
        def aplicar_anchos(row):
            for i, cell in enumerate(row.cells):
                cell.width = anchos[i]
        
        for i, width in enumerate(anchos):
            table.columns[i].width = width
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '9F2B20')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        aplicar_anchos(table.rows[0])
        
        # Filas
        for fila in filas:
            row_cells = table.add_row().cells
            for i, valor in enumerate(fila):
                row_cells[i].text = str(valor)
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            aplicar_anchos(table.rows[-1])
        
        return table

    def table_for_filesystem(headers, filas, ancho_total=Cm(17)):
        num_cols = len(headers)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        table.width = ancho_total
        
        # Definir anchos: 
        ancho_col1 = Cm(5)
        ancho_col6 = Cm(3)  

        
        columnas_fijas = {0: ancho_col1, 5: ancho_col6}
        columnas_restantes = num_cols - len(columnas_fijas)
        
        if columnas_restantes > 0:
            ancho_usado = sum(columnas_fijas.values())
            ancho_resto = int((ancho_total - ancho_usado) / columnas_restantes)
        
        anchos = []
        for i in range(num_cols):
            if i in columnas_fijas:
                anchos.append(columnas_fijas[i])
            else:
                anchos.append(ancho_resto)
        
        def aplicar_anchos(row):
            for i, cell in enumerate(row.cells):
                cell.width = anchos[i]
        
        for i, width in enumerate(anchos):
            table.columns[i].width = width
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '9F2B20')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        aplicar_anchos(table.rows[0])
        
        # Filas
        for fila in filas:
            row_cells = table.add_row().cells
            for i, valor in enumerate(fila):
                row_cells[i].text = str(valor)
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            aplicar_anchos(table.rows[-1])
        
        return table

    def table_for_backups(headers, filas, ancho_total=Cm(17)):
        num_cols = len(headers)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        table.width = ancho_total
        
        # Definir anchos: 
        ancho_col1 = Cm(2)
        ancho_col3 = Cm(3.24)
        ancho_col4 = Cm(2.5)  

        
        columnas_fijas = {0: ancho_col1, 2: ancho_col3,  3: ancho_col4}
        columnas_restantes = num_cols - len(columnas_fijas)
        
        if columnas_restantes > 0:
            ancho_usado = sum(columnas_fijas.values())
            ancho_resto = int((ancho_total - ancho_usado) / columnas_restantes)
        
        anchos = []
        for i in range(num_cols):
            if i in columnas_fijas:
                anchos.append(columnas_fijas[i])
            else:
                anchos.append(ancho_resto)
        
        def aplicar_anchos(row):
            for i, cell in enumerate(row.cells):
                cell.width = anchos[i]
        
        for i, width in enumerate(anchos):
            table.columns[i].width = width
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '9F2B20')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        aplicar_anchos(table.rows[0])
        
        # Filas
        for fila in filas:
            row_cells = table.add_row().cells
            for i, valor in enumerate(fila):
                row_cells[i].text = str(valor)
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            aplicar_anchos(table.rows[-1])
        
        return table
    
    def salto_pagina():
        doc.add_page_break()

    # =======================================================================================
    # =======================================================================================
    # =======================================================================================
    #                                   CONTENIDO DEL DOCUMENTO
    # =======================================================================================
    # =======================================================================================
    # =======================================================================================    
    # ---------- PORTADA (Página 1) ----------
    # Espacio
    for _ in range(1):
        parrafo_vacio(20)

    # Título pegado arriba, sin espacio antes, pero con espacio después para separar del subtítulo
    texto_posicionado(doc, "INFORME CONSOLIDADO", posicion='arriba', tamano=30, bold=True, color=(0,0,0))
    # ↑ por defecto space_before=0, space_after=0

    # Subtítulo pegado al título (0 espacio arriba y 0 abajo)
    texto_posicionado(doc, "ACTIVIDADES EFECTUADAS POR", posicion='arriba', tamano=20, color=(100,100,100))

    # NETGROUP pegado al subtítulo
    texto_posicionado(doc, "NETGROUP S.A.", posicion='arriba', tamano=25, bold=True, color=(0,0,0))

    # Espacio superior
    for _ in range(9):
        parrafo_vacio(20)
    
    # Nombre de la empresa (centro)
    nombre_empresa = empresa.get('nombre', 'EMPRESA NO IDENTIFICADA') if empresa else 'EMPRESA NO IDENTIFICADA'
    texto_centrado(nombre_empresa.upper(), tamano=28, bold=True, 
                   color=(0, 0, 0), espacio_despues=40)
    
    # Espacio
    for _ in range(3):
        parrafo_vacio(20)
    
    # ---------- DIRIGIDO A ----------
    dirigido_nombre = empresa.get('dirigido', '') if empresa else ''
    dirigido_cargo = empresa.get('cargo', '') if empresa else ''
    
    # Tabla especial para "Dirigido a"
    table_dir = doc.add_table(rows=0, cols=2)
    table_dir.autofit = False
    table_dir.allow_autofit = False
    table_dir.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    row = table_dir.add_row().cells
    row[0].text = "Dirigido a:"
    row[1].text = f"{dirigido_nombre}\n{dirigido_cargo}"
    
    row[0].width = Cm(4.5)
    row[1].width = Cm(10)
    
    # Formato celda 0
    for paragraph in row[0].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(80, 80, 80)
    
    # Formato celda 1 (nombre + cargo)
    for paragraph in row[1].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(64, 64, 64)
    
    # Quitar bordes
    for cell in row:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    
    
    # ---------- INFO ROWS ----------
    ip_servidor = conexion.get('ip_servidor', 'N/A') if conexion else 'N/A'
    fecha_informe = datos.get('fecha_informe', 'N/A')
    
    info_filas = [
        ["Servidor", ip_servidor],
        ["Recopilado por", "Área Bases de Datos"],
        ["Fecha del informe", fecha_informe.title() if isinstance(fecha_informe, str) else fecha_informe],
    ]
    
    tabla_sin_bordes(info_filas, [Cm(4.5), Cm(10)], alineacion_col1=WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Espacio antes del footer
    for _ in range(12):
        parrafo_vacio(0)
    
    # ---------- FOOTER ----------
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        "Este documento contiene información confidencial y es de uso "
        "exclusivo del cliente y del personal autorizado de NetGroup S.A."
    )
    footer_run.font.size = Pt(9)
    footer_run.font.name = 'Calibri'
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True

    # ---------- PÁGINA 2: CONTENIDOS ----------
    salto_pagina()
    
    contenidos = [
        ("1. Métricas Principales", "3"),
        ("2. Tablespaces", "4"),
        ("3. Resumen de Objetos", "5"),
        ("4. Estadísticas de Tablas", "5"),
        ("5. Sistema Operativo", "6"),
        ("6. Filesystems", "6"),
        ("7. Backups Recientes", "7"),
        ("8. Procesos Oracle", "8"),
    ]
    agregar_tabla_contenidos(doc, contenidos)
    
    # ---------- PÁGINA 3: MÉTRICAS ----------
    salto_pagina()

    agregar_titulo("1. Métricas Principales", nivel=1)

    # Espacio
    for _ in range(1):
        parrafo_vacio(20)
    
    metricas = [
        ["Tamaño BD", f"{datos.get('bd_tamano_tb', '0')} TB"],
        ["Objetos Inválidos", str(datos.get('obj_invalidos', 0))],
        ["Uso Filesystem", f"{datos.get('filesystem_pct', '0')}%"],
        ["RAM Usada", f"{datos.get('ram_usada_pct', '0')}%"],
    ]
    tabla_tipo_2(["", ""], metricas, mostrar_encabezado=False)

    # Espacio
    for _ in range(2):
        parrafo_vacio(20)

    agregar_titulo("Estado del servidor", nivel=1)

    # Espacio
    for _ in range(1):
        parrafo_vacio(20)

    estado_servidor = [
        ["Uptime", f"{datos.get('uptime', '0')}"],
        ["Load Average", str(datos.get('load_avg', 0))],
    ]
    tabla_tipo_2(["", ""],estado_servidor, mostrar_encabezado=False)

    # Espacio
    for _ in range(2):
        parrafo_vacio(20)

    agregar_titulo("Uso de Memoria", nivel=1)

    # Espacio
    for _ in range(1):
        parrafo_vacio(20)

    memoria_ram = [
        ["RAM Usada", f"{datos.get('ram_usada_gb', '0')} GB"],
        ["RAM Disponible", f"{datos.get('ram_disponible_gb', '0')} GB"],
        ["Total RAM",f"{datos.get('ram_total_gb', '0')} GB"],
    ]
    tabla_tipo_2(["", ""],memoria_ram, mostrar_encabezado=False)

    # Espacio
    for _ in range(2):
        parrafo_vacio(20)

    agregar_titulo("Procesos Claves de Oracle", nivel=1)

    # Espacio
    for _ in range(1):
        parrafo_vacio(20)

    procesos_ol =  datos.get('procesos_oracle_clave', [])

    procesos_oracle = []
    for proc in procesos_ol:
        procesos_oracle.append([
            proc.get('usuario', 'N/A'),
            proc.get('pid', '0'),
            proc.get('inicio', 'N/A'),
            proc.get('cpu_time', '0:00'),
            proc.get('comando', 'N/A'),
        ])

    agregar_tabla(["Usuario", "PID", "inicio", "CPU Time", "Proceso"],procesos_oracle)

    # ---------- PÁGINA 4: ESTADO DE LA INSTANCIA ----------
    salto_pagina()

    agregar_titulo("Estado de la Base de Datos", nivel=1)

    # Espacio
    for _ in range(1):
        parrafo_vacio(20)

    estado_instancia = datos.get('estado_instancia', [])
    
    estado_inst = [
        ["Current Status", str(estado_instancia.get('current_status', ''))],
        ["Up Since", str(estado_instancia.get('up_since', ''))],
        ["Instance Name", str(estado_instancia.get('instance_name', ''))],
        ["Database Version", str(estado_instancia.get('database_version', ''))],
        ["Database Status", str(estado_instancia.get('database_status', ''))],
        ["Shutdown Pending", str(estado_instancia.get('shutdown_pending', ''))],
        ["Active State", str(estado_instancia.get('active_state', ''))],
        ["Blocked", str(estado_instancia.get('blocked', ''))],
        ["Parallel", str(estado_instancia.get('parallel', ''))],
        ["Archiver", str(estado_instancia.get('archiver', ''))],
        ["Logins", str(estado_instancia.get('logins', ''))],
    ]
    tabla_tipo_2(["", ""], estado_inst, mostrar_encabezado=False)

    # Espacio
    for _ in range(2):
        parrafo_vacio(20)

    agregar_titulo("Rendimiento de la Instancia", nivel=1)

    # Espacio
    for _ in range(1):
        parrafo_vacio(20)

    rendimiento_inst = [
        ["Buffer Hit", f"{datos.get('buffer_hit', '0')} %"],
        ["System Global Area (SGA)", f"{datos.get('sga_gb', '0')} GB"],
        ["Program Global Area (PGA)", f"{datos.get('pga_gb', '0')} GB"],
    ]

    tabla_tipo_2(["",""], rendimiento_inst, mostrar_encabezado=False)

    # ---------- PÁGINA 5: TABLESPACES ----------
    salto_pagina()

    agregar_titulo("2. Tablespaces", nivel=1)

        # Espacio
    for _ in range(1):
        parrafo_vacio(20)
    
    ts_headers = ["Nombre", "File Name", "Asignado (MB)", "Usado (MB)", "% Uso", "Status", "Auto Extend"]
    ts_filas = []
    for ts in datos.get('tablespaces', []):
        ts_filas.append([
            ts.get('nombre', 'N/A'),
            ts.get('file_name', 'N/A'),
            ts.get('asignado_mb', 0),
            ts.get('usado_mb', 0),
            f"{ts.get('pct_uso', 0)}%",
            ts.get('status', 'N/A'),
            ts.get('autoextend', 'N/A'),
        ])
    table_for_tablespace(ts_headers, ts_filas)

    # ---------- PAGINA FILESYSTEMS ----------
    salto_pagina()

    agregar_titulo("Uso de File Systems", nivel=1)

        # Espacio
    for _ in range(1):
        parrafo_vacio(20)
    
    ts_headers = ["Filesystem", "Tamaño", "Usado", "Disponible", "Uso %", "Montado en"]
    filesystems = datos.get('filesystems', [])
    ts_filas = []

    for ts in filesystems:
        ts_filas.append([
            ts.get('filesystem', 'N/A'),
            ts.get('size', 0),
            ts.get('used', 0),
            ts.get('available', 0),
            ts.get('use_pct', 0),
            ts.get('mounted_on', 'N/A'),
        ])

    table_for_filesystem(ts_headers, ts_filas)

    # Espacio
    for _ in range(2):
        parrafo_vacio(20)

    agregar_titulo("Backups", nivel=1)

        # Espacio
    for _ in range(1):
        parrafo_vacio(20)
    
    ts_headers = ["Permisos", "Archivo", "Fecha", "Tamaño"]
    filesystems = datos.get('backups', [])
    ts_filas = []

    for ts in filesystems:
        ts_filas.append([
            ts.get('permisos', 'N/A'),
            ts.get('archivo', 'N/A'),
            ts.get('fecha', 'N/A'),
            filesizeformat(ts.get('size', 0))
        ])

    table_for_backups(ts_headers, ts_filas)
    
    # ---------- PÁGINA 4: OBJETOS Y STATS ----------
    salto_pagina()

    agregar_titulo("Resumen de Objetos", nivel=1)

    # Espacio
    for _ in range(1):
        parrafo_vacio(20)

    obj = datos.get('resumen_objetos', {})
    obj_filas = [
        ["Total Objetos", obj.get('total', 0)],
        ["Packages", obj.get('packages', 0)],
        ["Package Bodies", obj.get('package_bodies', 0)],
        ["Tablas", obj.get('tables', 0)],
        ["Funciones", obj.get('functions', 0)],
        ["Procedimientos", obj.get('procedures', 0)],
        ["Secuencias", obj.get('sequences', 0)],
        ["Triggers", obj.get('triggers', 0)],
        ["Índices", obj.get('indexes', 0)],
        ["Vistas", obj.get('views', 0)],
        ["Tipos", obj.get('types', 0)],
        ["Objetos Invalidos", datos.get('obj_invalidos', 0)],
        ["Errores de Compilacion", datos.get('errores_compilacion', 0)],             
        ["Ultima Modificacion", datos.get('ultima_modificacion_objeto', "N/A")],   
    ]
    tabla_tipo_2(["", ""], obj_filas, mostrar_encabezado=False)

    # Espacio
    for _ in range(2):
        parrafo_vacio(20)
    
    agregar_titulo("Estadísticas de Tablas", nivel=1)

    # Espacio
    for _ in range(1):
        parrafo_vacio(20)

    stats = datos.get('stats_tablas', {})
    tabla_tipo_2(
        ["", ""],
        [
            ["Actualizadas (≤7 días)", stats.get('actualizadas', 0)],
            ["No actualizadas (>7 días)", stats.get('no_actualizadas', 0)],
            ["Sin estadísticas", stats.get('sin_estadisticas', 0)],
        ],
        mostrar_encabezado=False
    )
      
    # ============================================================
    # GUARDAR (buffer o archivo)
    # ============================================================
    if buffer is not None:
        doc.save(buffer)
        return buffer
    else:
        doc.save(ruta_salida)
        return ruta_salida